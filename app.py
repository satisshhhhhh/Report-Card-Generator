from os import name
from docx import Document
import pandas as pd
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.shared import Pt
from docx2pdf import convert

# creating a word doc and converting it to PDF
def word_doc(fullName, nameOfSchool, dob, round, regNum):
    document = Document()
    print("Creating document...")
    document.add_picture('logo.png', width=Inches(1.0))
    document.add_heading(nameOfSchool, 0)

    table = document.add_table(rows=4, cols=2)
    cell = table.cell(0, 0)
    cell.text = 'Name: ' + fullName
    table.cell(0,1).text = 'School: ' + nameOfSchool
    row = table.rows[1]
    row.cells[0].text = (f'Grade: {grade}')
    row.cells[1].text = (f'Date Of Birth: {dob}')
    row = table.rows[2]
    row.cells[0].text = 'Country: ' + aStudent['Country_of_Residence'].values[0]
    row.cells[1].text = 'Gender: ' + aStudent['Gender'].values[0]
    row = table.rows[3]
    row.cells[0].text = 'City: ' + aStudent['City_of_Residence'].values[0]
    row.cells[1].text = 'Test Time & Date: ' + aStudent['Date_and_time_of_test'].values[0]

    p = document.add_paragraph(f'Round: {round}')
    p = document.add_paragraph(f'Registration Number: {regNum}')

    pic = "studentPics\\"
    png = ".png"
    picture = pic+fullName+png
    document.add_picture(picture, width=Inches(2.0))

    p = document.add_paragraph().add_run(f'Total Score scored: {totalScore} / {outOf}')
    p.font.size = Pt(20)

    p = document.add_paragraph().add_run(f'Graded {remark}')
    p.font.size = Pt(18)

    if (totalScore > 40):
        p = document.add_paragraph().add_run(f'            CONGRATULATIONS')
        p.font.size = Pt(32)
    else:
        p = document.add_paragraph().add_run(f'                  FAILED')
        p.font.size = Pt(32)

    p = document.add_paragraph().add_run(aStudent['Final_result'].values[0])
    p.font.size = Pt(15)

    #Please edit the path to your required path
    path = 'D:\My Projects\Report Card Generator Using Python\\'
    wordDoc = 'Report.docx'
    nameOfDoc = path+fullName+wordDoc
    document.save(nameOfDoc)
    print("Document Created")

    #Converting it to PDF
    print("Converting to PDF")
    convert(fullName+wordDoc)
    print("\nPDF Conversion Successful!")


#Using Pandas library to Read Excel File--
#Please edit the path to your required path
data = pd.read_excel("D:\My Projects\Report Card Generator Using Python\DummyData.xlsx", "Sheet1",  skiprows=1)
data.columns = [c.replace(' ', '_') for c in data.columns]

print("Registered Students:")
allStudents = data['Full_Name_'].unique() #All Students 
print(allStudents)
print("Please Enter Full Name of the student: ")
nameFromUser = input().upper()

#Checking for valid student name
while nameFromUser not in allStudents:
    if nameFromUser == 'Q':
        quit()
    else:
        print(f'Sorry, there is no student with name {nameFromUser}')
        print("PLease Try Again or press Q to exit:")
        nameFromUser = input().upper()

#Getting the score of the required student.
aStudent = data.loc[data["Full_Name_"] == nameFromUser]

#Getting student's score
totalScore = aStudent['Your_score'].sum()
#print(f'Total Score scored by {nameFromUser} : {totalScore}')

#Total score
outOf = aStudent['Score_if_correct'].sum()
#print(f'Total Scores are: {outOf}')

#Grade earned by the student
if totalScore >= 90:
    remark = "O"
elif totalScore >= 80 and totalScore <= 89.99:
    remark = "A+"
elif totalScore >= 70 and totalScore <= 79.99:
    remark = "A"
elif totalScore >= 60 and totalScore <= 69.99:
    remark = "B+"
elif totalScore >= 55 and totalScore <= 59.99:
    remark = "B"
elif totalScore >= 50 and totalScore <= 54.99:
    remark = "C"
elif totalScore >= 45 and totalScore <= 49.99:
    remark = "D"
elif totalScore >= 40 and totalScore <= 44.99:
    remark = "E"
else:
    remark = "F"
#print(f'{nameFromUser} is Graded {remark}')

#Getting the School of the student
nameOfSchool = aStudent['Name_of_School_'].values[0]
#print(f'{nameFromUser} used to go to school called {nameOfSchool}')

#Getting the Round of the student
round = ''
round = aStudent['Round'].values[0]
#print(round)

grade = ''
grade = aStudent['Grade_'].values[0]

dob = ''
ts = pd.to_datetime(str(aStudent['Date_of_Birth_'].values[0])) 
dob = ts.strftime('%d-%m-%Y')
#print(dob)

regNum = ''
regNum = aStudent['Registration_Number'].values[0]

word_doc(nameFromUser, nameOfSchool, dob, round, regNum)
